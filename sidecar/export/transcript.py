"""Exportação da transcrição (por áudio e por projeto).

Monta uma representação intermediária única — linhas de segmento com tempo,
falante, texto e (opcionalmente) os códigos qualitativos — e a serializa em cada
formato. Dados vêm do `repository` (fonte única). Importações de DOCX/PDF são
adiadas para dentro dos formatadores que as usam (pesadas, opcionais).
"""
from __future__ import annotations

from typing import Any

import repository as repo

from .common import (
    MEDIA_TYPES,
    Export,
    TableExport,
    rows_to_csv,
    rows_to_json,
    slugify,
    srt_time,
    vtt_time,
)

# Colunas dos formatos tabulares. `codes` só aparece quando include_coding.
_BASE_COLUMNS = ["audio_id", "filename", "seq", "start", "end", "speaker", "text"]


def _segment_rows(
    audio_id: int, filename: str, *, coding: dict[int, list[dict[str, Any]]] | None
) -> list[dict[str, Any]]:
    """Linhas planas de um áudio (uma por segmento), prontas para tabela/JSON."""
    rows: list[dict[str, Any]] = []
    for seg in repo.list_segments(audio_id):
        row: dict[str, Any] = {
            "audio_id": audio_id,
            "filename": filename,
            "seq": seg["seq"],
            "start": seg["start"],
            "end": seg["end"],
            "speaker": seg["speaker"],
            "text": seg["text"],
        }
        if coding is not None:
            codes = coding.get(seg["id"], [])
            row["codes"] = "; ".join(c["name"] for c in codes)
        rows.append(row)
    return rows


def _collect(
    *, project_id: int | None, audio_id: int | None, include_coding: bool
) -> tuple[str, list[dict[str, Any]]]:
    """Reúne as linhas do escopo pedido e devolve `(titulo, linhas)`.

    Para um projeto, concatena os áudios transcritos na ordem de criação. Para um
    áudio, apenas ele. `include_coding` agrega os códigos por segmento.
    """
    if audio_id is not None:
        audio = repo.get_audio(audio_id)
        title = audio["filename"] if audio else f"audio-{audio_id}"
        coding = repo.coding_for_audio(audio_id) if include_coding else None
        return title, _segment_rows(audio_id, title, coding=coding)

    assert project_id is not None
    project = repo.get_project(project_id)
    title = project["name"] if project else f"projeto-{project_id}"
    rows: list[dict[str, Any]] = []
    for audio in repo.list_audios(project_id):
        if audio["status"] != "done":
            continue
        coding = repo.coding_for_audio(audio["id"]) if include_coding else None
        rows.extend(_segment_rows(audio["id"], audio["filename"], coding=coding))
    return title, rows


def _columns(include_coding: bool) -> list[str]:
    return [*_BASE_COLUMNS, "codes"] if include_coding else _BASE_COLUMNS


def _format_csv(
    title: str, rows: list[dict[str, Any]], include_coding: bool, *, delimiter: str
) -> bytes:
    return rows_to_csv(
        TableExport(_columns(include_coding), rows), delimiter=delimiter
    )


def _format_json(title: str, rows: list[dict[str, Any]], include_coding: bool) -> bytes:
    return rows_to_json({"title": title, "segments": rows})


def _format_subtitle(rows: list[dict[str, Any]], fmt: str) -> bytes:
    """SRT ou VTT a partir das linhas de segmento (numeração sequencial global)."""
    time = vtt_time if fmt == "vtt" else srt_time
    blocks: list[str] = ["WEBVTT", ""] if fmt == "vtt" else []
    for i, r in enumerate(rows, start=1):
        speaker = f"{r['speaker']}: " if r["speaker"] else ""
        blocks.append(str(i))
        blocks.append(f"{time(r['start'])} --> {time(r['end'])}")
        blocks.append(f"{speaker}{r['text']}".strip())
        blocks.append("")
    return "\n".join(blocks).encode("utf-8")


def _format_docx(title: str, rows: list[dict[str, Any]], include_coding: bool) -> bytes:
    import io

    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    for r in rows:
        stamp = f"[{srt_time(r['start'])}]"
        who = f" {r['speaker']}" if r["speaker"] else ""
        p = doc.add_paragraph()
        p.add_run(f"{stamp}{who}").bold = True
        p.add_run(f"  {r['text']}")
        if include_coding and r.get("codes"):
            tag = p.add_run(f"  ⟦{r['codes']}⟧")
            tag.italic = True
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _format_pdf(title: str, rows: list[dict[str, Any]], include_coding: bool) -> bytes:
    import io

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    from xml.sax.saxutils import escape

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, title=title)
    styles = getSampleStyleSheet()
    flow = [Paragraph(escape(title), styles["Title"]), Spacer(1, 12)]
    for r in rows:
        stamp = f"[{srt_time(r['start'])}]"
        who = f" {escape(r['speaker'])}" if r["speaker"] else ""
        codes = (
            f' <i>⟦{escape(r["codes"])}⟧</i>'
            if include_coding and r.get("codes")
            else ""
        )
        flow.append(
            Paragraph(
                f"<b>{stamp}{who}</b>  {escape(r['text'])}{codes}", styles["BodyText"]
            )
        )
    doc.build(flow)
    return buf.getvalue()


def export_transcript(
    fmt: str,
    *,
    project_id: int | None = None,
    audio_id: int | None = None,
    include_coding: bool = False,
) -> Export:
    """Exporta a transcrição de um áudio ou projeto no formato pedido.

    `fmt`: csv | tsv | json | srt | vtt | docx | pdf. `include_coding` adiciona os
    códigos qualitativos (ignorado em SRT/VTT, que são só legenda).
    """
    title, rows = _collect(
        project_id=project_id, audio_id=audio_id, include_coding=include_coding
    )
    base = slugify(title)

    if fmt == "csv":
        content = _format_csv(title, rows, include_coding, delimiter=",")
    elif fmt == "tsv":
        content = _format_csv(title, rows, include_coding, delimiter="\t")
    elif fmt == "json":
        content = _format_json(title, rows, include_coding)
    elif fmt in ("srt", "vtt"):
        content = _format_subtitle(rows, fmt)
    elif fmt == "docx":
        content = _format_docx(title, rows, include_coding)
    elif fmt == "pdf":
        content = _format_pdf(title, rows, include_coding)
    else:
        raise ValueError(f"Formato de exportação não suportado: {fmt}")

    return Export(content, MEDIA_TYPES[fmt], f"{base}.{fmt}")
